"""
Deterministic workflow execution engine.

A workflow definition is a DAG: {"nodes": [...], "edges": [...]}.
Each node has {id, type, data: {label, config}}. Edges may carry a
sourceHandle ("true"/"false") for condition branching.

The engine executes nodes in topological order, persisting a
WorkflowNodeRun row per node (status pending → running → succeeded/
failed/skipped) so the UI can poll live activity.

Template syntax inside node config values:
    {{trigger.someField}}         — value from trigger input
    {{node_id.output.path.0.x}}   — output of an upstream node
If a string is exactly one template, the raw value (dict/list/number)
is passed through; otherwise values are interpolated as strings.
"""
import asyncio
import json
import logging
import re
import os
from datetime import datetime, timezone
from typing import Any, Callable, Dict, List, Optional

import requests as http_requests
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.workflow import Workflow, WorkflowRun, WorkflowNodeRun
from app.models.document import Document
from app.models.job import Job
from app.models.integration import Integration, IntegrationType, IntegrationStatus
from app.models.ai_settings import AISettings
from app.utils.redact import redact_secrets

logger = logging.getLogger(__name__)

# Storage-service key prefix for write_output files. With local storage
# (base /app/uploads) this resolves to /app/uploads/workflow_outputs — the
# same shared-volume location the engine wrote to before; with MinIO/S3 the
# files land in the bucket so any replica can serve them.
WORKFLOW_OUTPUT_DIR = os.environ.get("WORKFLOW_OUTPUT_DIR", "workflow_outputs")

# ── Node catalog (exposed to the frontend palette) ──────────────────
NODE_TYPES: List[Dict[str, Any]] = [
    {
        "type": "trigger_manual",
        "category": "trigger",
        "label": "Manual Trigger",
        "description": "เริ่ม workflow ด้วยตนเอง (กดปุ่ม Run) พร้อมส่ง input JSON ได้",
        "config_fields": [],
        "output_fields": [],  # ฟิลด์ขึ้นกับ input JSON ที่ผู้ใช้ส่งตอน Run
    },
    {
        "type": "trigger_schedule",
        "category": "trigger",
        "label": "Schedule Trigger",
        "description": "เริ่ม workflow ตามตารางเวลา (cron) ที่ตั้งไว้ใน workflow settings",
        "config_fields": [],
        "output_fields": [{"name": "scheduled_at", "label": "เวลาที่ทริกเกอร์"}],
    },
    {
        "type": "trigger_webhook",
        "category": "trigger",
        "label": "Webhook Trigger",
        "description": "เริ่ม workflow จาก webhook ภายนอก เช่น web application หรือ LINE webhook",
        "config_fields": [],
        "output_fields": [
            {"name": "body", "label": "Payload body"},
            {"name": "query", "label": "Query parameters"},
            {"name": "headers", "label": "Request headers"},
            {"name": "method", "label": "HTTP method"},
            {"name": "received_at", "label": "เวลาที่รับ webhook"},
        ],
    },
    {
        "type": "job_source",
        "category": "data",
        "label": "Jobs",
        "description": "นำข้อมูลที่ประมวลผลแล้วจาก Job (extracted/reviewed data) เข้าสู่ workflow",
        "config_fields": [
            {"name": "job_id", "label": "เลือก Job", "type": "job_select", "required": True,
             "hint": "เลือก Job ที่มีเอกสารประมวลผลแล้ว"},
            {"name": "data_source", "label": "ข้อมูลที่ใช้", "type": "select",
             "options": ["reviewed", "extracted", "ocr_text"], "default": "reviewed",
             "hint": "reviewed = ข้อมูลที่ตรวจแล้ว (แนะนำ), ocr_text = ข้อความดิบ"},
            {"name": "status", "label": "กรองตามสถานะเอกสาร", "type": "select",
             "options": ["", "extraction_completed", "reviewed", "ocr_completed"], "required": False,
             "hint": "เว้นว่าง = ไม่กรองเพิ่ม"},
            {"name": "only_completed", "label": "เฉพาะเอกสารที่ประมวลผลเสร็จ", "type": "boolean", "default": True},
            {"name": "limit", "label": "จำนวนเอกสารสูงสุด", "type": "number", "default": 50,
             "placeholder": "50"},
        ],
        "output_fields": [
            {"name": "count", "label": "จำนวนเอกสาร"},
            {"name": "records", "label": "ข้อมูลทั้งหมด (array)"},
            {"name": "documents", "label": "รายการเอกสาร"},
            {"name": "job_name", "label": "ชื่อ Job"},
            {"name": "job_status", "label": "สถานะ Job"},
        ],
    },
    {
        "type": "document_source",
        "category": "data",
        "label": "Document Source",
        "description": "ดึงเอกสาร (OCR text + extracted data) จาก Job ที่เลือก",
        "config_fields": [
            {"name": "job_id", "label": "เลือก Job", "type": "job_select", "required": True},
            {"name": "status", "label": "กรองตามสถานะเอกสาร", "type": "select",
             "options": ["", "extraction_completed", "reviewed", "ocr_completed"], "required": False,
             "hint": "เว้นว่าง = ไม่กรองเพิ่ม"},
            {"name": "limit", "label": "จำนวนเอกสารสูงสุด", "type": "number", "default": 10,
             "placeholder": "10"},
            {"name": "include_ocr_text", "label": "รวมข้อความ OCR", "type": "boolean", "default": True},
        ],
        "output_fields": [
            {"name": "count", "label": "จำนวนเอกสาร"},
            {"name": "documents", "label": "รายการเอกสาร (มี ocr_text)"},
            {"name": "job_name", "label": "ชื่อ Job"},
        ],
    },
    {
        "type": "llm",
        "category": "ai",
        "label": "LLM / Agent",
        "description": "ส่งข้อมูลให้ LLM วิเคราะห์ สรุป แปลง หรือตรวจสอบตาม prompt",
        "config_fields": [
            {"name": "ai_provider_id", "label": "AI Agent Provider", "type": "ai_provider_select", "required": False,
             "hint": "เลือก provider จาก Setting AI; เว้นว่าง = ใช้ Agent Provider ที่ตั้งไว้กลางระบบ"},
            {"name": "system_prompt", "label": "System prompt", "type": "textarea", "required": False,
             "placeholder": "คุณเป็นผู้ช่วยสรุปข้อมูลเอกสาร ตอบเป็นภาษาไทย กระชับ",
             "hint": "กำหนดบทบาท/สไตล์การตอบของ AI"},
            {"name": "prompt", "label": "Prompt", "type": "textarea", "required": True,
             "placeholder": "สรุปรายการต่อไปนี้เป็น bullet:\n\n{{job_source_xxx.records}}",
             "hint": "ใช้ปุ่ม “+ แทรกข้อมูล” ด้านบนช่องเพื่ออ้างผลจากโหนดก่อนหน้า"},
            {"name": "json_output", "label": "แปลงคำตอบเป็น JSON", "type": "boolean", "default": False,
             "hint": "เปิดเมื่อสั่งให้ AI ตอบเป็น JSON แล้วต้องการใช้ฟิลด์ data ต่อ"},
        ],
        "output_fields": [
            {"name": "text", "label": "ข้อความตอบกลับ"},
            {"name": "data", "label": "JSON ที่ parse แล้ว (ถ้าเปิด)"},
        ],
    },
    {
        "type": "condition",
        "category": "logic",
        "label": "Condition (If/Else)",
        "description": "ตรวจเงื่อนไขแล้วแยกเส้นทาง True / False",
        "config_fields": [
            {"name": "left", "label": "ค่าที่ตรวจ", "type": "text", "required": True,
             "placeholder": "{{job_source_xxx.count}}",
             "hint": "ใช้ปุ่ม “+ แทรกข้อมูล” เพื่อเลือกค่าจากโหนดก่อนหน้า"},
            {"name": "operator", "label": "เงื่อนไข", "type": "select",
             "options": ["equals", "not_equals", "contains", "not_contains", "greater_than",
                         "less_than", "is_empty", "is_not_empty"], "default": "equals"},
            {"name": "right", "label": "ค่าที่ใช้เทียบ", "type": "text", "required": False,
             "placeholder": "0",
             "hint": "ค่าที่ใช้เทียบ เช่น 0, reviewed (ไม่ต้องใส่ถ้าใช้ is_empty/is_not_empty)"},
        ],
        "output_fields": [{"name": "result", "label": "ผลลัพธ์ true/false"}],
    },
    {
        "type": "transform",
        "category": "logic",
        "label": "Transform / Mapping",
        "description": "สร้าง object ใหม่จากการ map ค่าด้วย template",
        "config_fields": [
            {"name": "mappings", "label": "การ map ฟิลด์", "type": "mappings", "required": True,
             "hint": "ตั้งชื่อฟิลด์ใหม่ทางซ้าย แล้วใช้ปุ่มแทรกข้อมูลเลือกค่าทางขวา"},
        ],
        "output_fields": [],  # ฟิลด์ขึ้นกับ target ที่ผู้ใช้กำหนด (เติมแบบไดนามิกฝั่ง UI)
    },
    {
        "type": "python_code",
        "category": "developer",
        "label": "Python Code",
        "description": "รันโค้ด Python ใน sandbox ปลอดภัย — อ่านข้อมูลจาก inputs แล้วเซ็ตตัวแปร result",
        "config_fields": [
            {"name": "code", "label": "โค้ด Python", "type": "code", "required": True,
             "hint": "อ่านข้อมูลจากตัวแปร inputs แล้วเซ็ตตัวแปร result เป็นผลลัพธ์"},
            {"name": "input", "label": "Input", "type": "textarea", "required": False,
             "placeholder": "{{transform_xxx}}",
             "hint": "ค่านี้จะกลายเป็นตัวแปร inputs ในโค้ด"},
            {"name": "timeout", "label": "Timeout (วินาที)", "type": "number", "default": 30,
             "placeholder": "30"},
        ],
        "output_fields": [
            {"name": "result", "label": "ผลลัพธ์ (ตัวแปร result)"},
            {"name": "stdout", "label": "ข้อความที่ print"},
        ],
    },
    {
        "type": "http_request",
        "category": "action",
        "label": "HTTP Request",
        "description": "ส่งข้อมูลต่อไปยังระบบอื่นผ่าน REST API / Webhook",
        "config_fields": [
            {"name": "method", "label": "Method", "type": "select",
             "options": ["POST", "GET", "PUT", "PATCH", "DELETE"], "default": "POST"},
            {"name": "url", "label": "URL", "type": "text", "required": True,
             "placeholder": "https://example.com/webhook"},
            {"name": "headers", "label": "Headers (JSON)", "type": "textarea", "required": False,
             "placeholder": '{ "Content-Type": "application/json" }'},
            {"name": "body", "label": "Body", "type": "textarea", "required": False,
             "placeholder": "{{transform_xxx}}",
             "hint": "ใช้ปุ่มแทรกข้อมูลเพื่อส่งผลจากโหนดก่อนหน้า"},
        ],
        "output_fields": [
            {"name": "status_code", "label": "HTTP status"},
            {"name": "body", "label": "เนื้อหาที่ตอบกลับ"},
        ],
    },
    {
        "type": "write_output",
        "category": "action",
        "label": "Write Output",
        "description": "เขียนผลลัพธ์เป็นไฟล์ (JSON / Text / CSV / Excel / Word) เพื่อนำไปใช้ต่อ",
        "config_fields": [
            {"name": "filename", "label": "ชื่อไฟล์", "type": "text", "default": "output.json",
             "placeholder": "report.json",
             "hint": "นามสกุลไฟล์จะปรับให้ตรงกับรูปแบบที่เลือกโดยอัตโนมัติ"},
            {"name": "format", "label": "รูปแบบไฟล์", "type": "select",
             "options": ["json", "text", "csv", "xlsx", "docx"], "default": "json",
             "hint": "xlsx/docx: ถ้าเนื้อหาเป็นรายการ object (เช่น records) จะสร้างเป็นตาราง"},
            {"name": "content", "label": "เนื้อหา", "type": "textarea", "required": True,
             "placeholder": "{{transform_xxx}}",
             "hint": "ใช้ปุ่มแทรกข้อมูลเลือกผลลัพธ์ที่ต้องการบันทึก"},
        ],
        "output_fields": [
            {"name": "filename", "label": "ชื่อไฟล์"},
            {"name": "size", "label": "ขนาด (ตัวอักษร)"},
            {"name": "preview", "label": "ตัวอย่างเนื้อหา"},
        ],
    },
    {
        "type": "webhook_response",
        "category": "action",
        "label": "Webhook Response",
        "description": "กำหนด result ที่ caller จะอ่านได้จาก webhook poll endpoint",
        "config_fields": [
            {"name": "visible", "label": "ใช้เป็น result ของ webhook", "type": "boolean", "default": True},
            {"name": "status_code", "label": "HTTP status", "type": "number", "default": 200,
             "placeholder": "200"},
            {"name": "body", "label": "Result body", "type": "textarea", "required": True,
             "placeholder": "{{llm_1.text}}",
             "hint": "ใช้ template เพื่อเลือกผลจากโหนดก่อนหน้า เช่น {{trigger.body.events.0.message.text}}"},
            {"name": "condition_left", "label": "เงื่อนไข: ค่าที่ตรวจ", "type": "text", "required": False,
             "placeholder": "{{condition_1.result}}"},
            {"name": "condition_operator", "label": "เงื่อนไข", "type": "select",
             "options": ["", "equals", "not_equals", "contains", "not_contains", "greater_than",
                         "less_than", "is_empty", "is_not_empty"], "default": ""},
            {"name": "condition_right", "label": "เงื่อนไข: ค่าที่ใช้เทียบ", "type": "text", "required": False,
             "placeholder": "true"},
        ],
        "output_fields": [
            {"name": "visible", "label": "แสดงผลหรือไม่"},
            {"name": "status_code", "label": "HTTP status"},
            {"name": "body", "label": "Result body"},
        ],
    },
    {
        "type": "gdrive_upload",
        "category": "storage",
        "label": "Google Drive: อัปโหลด",
        "description": "อัปโหลดผลลัพธ์ของ workflow ขึ้นโฟลเดอร์ Google Drive",
        "config_fields": [
            {"name": "integration_id", "label": "บัญชี Google Drive", "type": "integration_select",
             "provider": "gdrive", "required": True,
             "hint": "เลือก credential ที่สร้างไว้ในเมนู Integration (ชนิด Google Drive)"},
            {"name": "folder_id", "label": "Folder ID ปลายทาง", "type": "text", "required": True,
             "placeholder": "1AbC...xyz",
             "hint": "คัดลอกจาก URL ของโฟลเดอร์ Drive และต้องแชร์โฟลเดอร์ให้อีเมล service account"},
            {"name": "filename", "label": "ชื่อไฟล์", "type": "text", "default": "result.json",
             "placeholder": "result.json"},
            {"name": "mime_type", "label": "ชนิดไฟล์ (MIME)", "type": "text", "default": "application/json",
             "placeholder": "application/json"},
            {"name": "content", "label": "เนื้อหา", "type": "textarea", "required": True,
             "placeholder": "{{transform_xxx}}",
             "hint": "ใช้ปุ่มแทรกข้อมูลเลือกผลลัพธ์ที่ต้องการอัปโหลด"},
        ],
        "output_fields": [
            {"name": "file_id", "label": "Drive file id"},
            {"name": "name", "label": "ชื่อไฟล์"},
            {"name": "link", "label": "ลิงก์เปิดไฟล์"},
        ],
    },
    {
        "type": "gdrive_import",
        "category": "storage",
        "label": "Google Drive: นำเข้า Job",
        "description": "ดึงทุกไฟล์จากโฟลเดอร์ Google Drive เข้า Job แล้วประมวลผล (OCR) ตามฟังก์ชัน Jobs",
        "config_fields": [
            {"name": "integration_id", "label": "บัญชี Google Drive", "type": "integration_select",
             "provider": "gdrive", "required": True},
            {"name": "folder_id", "label": "Folder ID ต้นทาง", "type": "text", "required": True,
             "placeholder": "1AbC...xyz",
             "hint": "ต้องแชร์โฟลเดอร์ให้อีเมล service account (สิทธิ์อ่าน)"},
            {"name": "job_id", "label": "นำเข้าไปยัง Job", "type": "job_select", "required": True},
            {"name": "name_filter", "label": "กรองชื่อไฟล์ (optional)", "type": "text", "required": False,
             "placeholder": ".pdf",
             "hint": "เว้นว่าง = ทุกไฟล์; ใส่นามสกุล/คำเช่น .pdf เพื่อกรอง"},
            {"name": "limit", "label": "จำนวนไฟล์สูงสุด", "type": "number", "default": 20, "placeholder": "20"},
        ],
        "output_fields": [
            {"name": "count", "label": "จำนวนไฟล์ที่นำเข้า"},
            {"name": "imported", "label": "รายการที่นำเข้า"},
            {"name": "job_id", "label": "Job ปลายทาง"},
        ],
    },
    {
        "type": "onedrive_upload",
        "category": "storage",
        "label": "OneDrive: อัปโหลด",
        "description": "อัปโหลดผลลัพธ์ของ workflow ขึ้นโฟลเดอร์ OneDrive / SharePoint",
        "config_fields": [
            {"name": "integration_id", "label": "บัญชี OneDrive", "type": "integration_select",
             "provider": "onedrive", "required": True,
             "hint": "เลือก credential ที่สร้างไว้ในเมนู Integration (ชนิด OneDrive)"},
            {"name": "folder_id", "label": "Folder item id (เว้นว่าง = root)", "type": "text", "required": False,
             "placeholder": "root",
             "hint": "ระบุ item id ของโฟลเดอร์ หรือเว้นว่างเพื่อใช้รากของ drive"},
            {"name": "filename", "label": "ชื่อไฟล์", "type": "text", "default": "result.json",
             "placeholder": "result.json"},
            {"name": "mime_type", "label": "ชนิดไฟล์ (MIME)", "type": "text", "default": "application/json",
             "placeholder": "application/json"},
            {"name": "content", "label": "เนื้อหา", "type": "textarea", "required": True,
             "placeholder": "{{transform_xxx}}",
             "hint": "ใช้ปุ่มแทรกข้อมูลเลือกผลลัพธ์ที่ต้องการอัปโหลด (ไฟล์ ≤4MB)"},
        ],
        "output_fields": [
            {"name": "file_id", "label": "OneDrive item id"},
            {"name": "name", "label": "ชื่อไฟล์"},
            {"name": "link", "label": "ลิงก์เปิดไฟล์"},
        ],
    },
    {
        "type": "onedrive_import",
        "category": "storage",
        "label": "OneDrive: นำเข้า Job",
        "description": "ดึงทุกไฟล์จากโฟลเดอร์ OneDrive / SharePoint เข้า Job แล้วประมวลผล (OCR)",
        "config_fields": [
            {"name": "integration_id", "label": "บัญชี OneDrive", "type": "integration_select",
             "provider": "onedrive", "required": True},
            {"name": "folder_id", "label": "Folder item id (เว้นว่าง = root)", "type": "text", "required": False,
             "placeholder": "root"},
            {"name": "job_id", "label": "นำเข้าไปยัง Job", "type": "job_select", "required": True},
            {"name": "name_filter", "label": "กรองชื่อไฟล์ (optional)", "type": "text", "required": False,
             "placeholder": ".pdf",
             "hint": "เว้นว่าง = ทุกไฟล์; ใส่นามสกุล/คำเช่น .pdf เพื่อกรอง"},
            {"name": "limit", "label": "จำนวนไฟล์สูงสุด", "type": "number", "default": 20, "placeholder": "20"},
        ],
        "output_fields": [
            {"name": "count", "label": "จำนวนไฟล์ที่นำเข้า"},
            {"name": "imported", "label": "รายการที่นำเข้า"},
            {"name": "job_id", "label": "Job ปลายทาง"},
        ],
    },
]

TEMPLATE_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_\-\.฀-๿]+)\s*\}\}")


# ── Template resolution ──────────────────────────────────────────────
def _lookup_path(context: Dict[str, Any], path: str) -> Any:
    parts = path.split(".")
    cur: Any = context
    for part in parts:
        if isinstance(cur, dict):
            if part in cur:
                cur = cur[part]
                continue
            return None
        if isinstance(cur, list):
            try:
                cur = cur[int(part)]
                continue
            except (ValueError, IndexError):
                return None
        return None
    return cur


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


def resolve_template(value: Any, context: Dict[str, Any]) -> Any:
    """Resolve {{path}} templates in a string / dict / list recursively."""
    if isinstance(value, str):
        match = TEMPLATE_RE.fullmatch(value.strip())
        if match:
            return _lookup_path(context, match.group(1))
        return TEMPLATE_RE.sub(lambda m: _stringify(_lookup_path(context, m.group(1))), value)
    if isinstance(value, dict):
        return {k: resolve_template(v, context) for k, v in value.items()}
    if isinstance(value, list):
        return [resolve_template(v, context) for v in value]
    return value


# ── Node executors ───────────────────────────────────────────────────
class NodeExecutionError(Exception):
    pass


def _exec_trigger(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    log("Workflow triggered")
    return context.get("trigger") or {}


def _exec_document_source(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    job_id = config.get("job_id")
    if not job_id:
        raise NodeExecutionError("Document Source: job_id is required")
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise NodeExecutionError(f"Job not found: {job_id}")

    query = db.query(Document).filter(Document.job_id == job_id)
    status = config.get("status")
    if status:
        query = query.filter(Document.status == status)
    limit = int(config.get("limit") or 10)
    docs = query.order_by(Document.uploaded_at.desc()).limit(limit).all()
    include_ocr = config.get("include_ocr_text", True)

    log(f"Job '{job.name or job_id}': loaded {len(docs)} document(s)")
    documents = []
    for d in docs:
        item: Dict[str, Any] = {
            "id": str(d.id),
            "filename": d.filename,
            "status": d.status,
            "extracted_data": d.reviewed_data or d.extracted_data,
        }
        if include_ocr:
            item["ocr_text"] = d.ocr_text
        documents.append(item)
    return {"job_id": str(job_id), "job_name": job.name, "count": len(documents), "documents": documents}


COMPLETED_DOC_STATUSES = {"extraction_completed", "reviewed"}


def _exec_job_source(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    """Bring processed data from a Job into the workflow.

    Output:
        {job_id, job_name, job_status, count,
         records: [<data per document>],   # convenient for downstream LLM/Transform
         documents: [{id, filename, status, data}]}
    """
    job_id = config.get("job_id")
    if not job_id:
        raise NodeExecutionError("Jobs node: job_id is required")
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise NodeExecutionError(f"Job not found: {job_id}")

    query = db.query(Document).filter(Document.job_id == job_id)
    status = config.get("status")
    if status:
        query = query.filter(Document.status == status)
    elif config.get("only_completed", True):
        query = query.filter(Document.status.in_(COMPLETED_DOC_STATUSES))

    limit = int(config.get("limit") or 50)
    docs = query.order_by(Document.uploaded_at.desc()).limit(limit).all()

    data_source = (config.get("data_source") or "reviewed").lower()
    log(f"Job '{job.name or job_id}' (status={job.status}): loaded {len(docs)} document(s) [data_source={data_source}]")

    documents: List[Dict[str, Any]] = []
    records: List[Any] = []
    for d in docs:
        if data_source == "ocr_text":
            data: Any = d.ocr_text
        elif data_source == "extracted":
            data = d.extracted_data
        else:  # reviewed — prefer reviewed_data, fall back to extracted_data
            data = d.reviewed_data if d.reviewed_data is not None else d.extracted_data
        records.append(data)
        documents.append({
            "id": str(d.id),
            "filename": d.filename,
            "status": d.status,
            "data": data,
        })

    return {
        "job_id": str(job_id),
        "job_name": job.name,
        "job_status": job.status,
        "count": len(documents),
        "records": records,
        "documents": documents,
    }


def _normalize_openai_base_url(base_url: Optional[str]) -> Optional[str]:
    if not base_url:
        return None
    normalized = base_url.strip().rstrip("/")
    if normalized.lower().endswith("/chat/completions"):
        normalized = normalized[: -len("/chat/completions")]
    return normalized


def _ai_setting_provider(setting: AISettings, model: Optional[str], source: str) -> Dict[str, Any]:
    provider_type = getattr(setting, "provider_type", None) or "completion_messages"
    resolved_model = (model or "").strip() or getattr(setting, "model", None) or "gpt-4o-mini"
    if provider_type == "openai_compatible":
        return {
            "provider": "openai_compatible",
            "apiKey": setting.api_key,
            "baseUrl": setting.api_url,
            "model": resolved_model,
            "source": source,
            "name": setting.display_name or setting.name,
        }
    return {
        "provider": "completion_messages",
        "apiUrl": setting.api_url,
        "apiKey": setting.api_key,
        "model": resolved_model,
        "source": source,
        "name": setting.display_name or setting.name,
    }


def _ensure_integration_owner(integration: Integration, owner_user_id: Optional[str]) -> None:
    """Block a workflow from using another user's Integration credentials."""
    if owner_user_id is None:
        return  # legacy workflow without an owner — nothing to enforce
    if integration.user_id is not None and str(integration.user_id) != str(owner_user_id):
        raise NodeExecutionError(
            f"Integration '{integration.name}' belongs to another user"
        )


def resolve_llm_provider(
    db: Session,
    integration_id: Optional[str] = None,
    model: Optional[str] = None,
    log: Optional[Callable[[str], None]] = None,
    ai_provider_id: Optional[str] = None,
    owner_user_id: Optional[str] = None,
) -> Dict[str, Any]:
    """Resolve the provider for Workflow LLM nodes.

    Priority:
    1. Explicit AI Settings provider selected on the node
    2. Legacy explicit LLM Integration ID on the node
    3. Setting AI > Agent Provider
    4. System AGENT_PROVIDER_* env fallback
    5. Default active AI Setting / OPENAI_API_KEY fallback
    6. First active LLM Integration for backward compatibility
    """

    def _log(msg: str) -> None:
        if log:
            log(msg)

    requested_model = (model or "").strip() or None

    if ai_provider_id:
        ai_setting = db.query(AISettings).filter(AISettings.id == ai_provider_id, AISettings.is_active == True).first()
        if not ai_setting:
            raise NodeExecutionError(f"AI provider not found or inactive: {ai_provider_id}")
        if not ai_setting.api_url or not ai_setting.api_key:
            raise NodeExecutionError(f"AI provider '{ai_setting.display_name or ai_setting.name}' is missing URL or key")
        provider = _ai_setting_provider(ai_setting, requested_model, "workflow_ai_provider")
        _log(f"Using selected AI provider '{provider['name']}' ({provider['provider']}, model={provider['model']})")
        return provider

    if integration_id:
        integration = db.query(Integration).filter(Integration.id == integration_id).first()
        if not integration:
            raise NodeExecutionError(f"Integration not found: {integration_id}")
        _ensure_integration_owner(integration, owner_user_id)
        icfg = integration.config or {}
        api_key = icfg.get("apiKey")
        if not api_key:
            raise NodeExecutionError(f"LLM integration '{integration.name}' is missing apiKey")
        provider = {
            "provider": "openai_compatible",
            "apiKey": api_key,
            "baseUrl": icfg.get("baseUrl") or None,
            "model": requested_model or icfg.get("model") or "gpt-4o-mini",
            "source": "workflow_llm_integration",
            "name": integration.name,
        }
        _log(f"Using LLM integration '{integration.name}' (model={provider['model']})")
        return provider

    agent_setting = (
        db.query(AISettings)
        .filter(AISettings.is_agent_provider == True, AISettings.is_active == True)
        .first()
    )
    if agent_setting and agent_setting.api_url and agent_setting.api_key:
        provider = _ai_setting_provider(agent_setting, requested_model, "ai_settings_agent_provider")
        _log(f"Using AI Settings Agent Provider '{provider['name']}' ({provider['provider']}, model={provider['model']})")
        return provider

    if settings.AGENT_PROVIDER_KEY:
        provider = {
            "provider": "openai_compatible",
            "apiKey": settings.AGENT_PROVIDER_KEY,
            "baseUrl": settings.AGENT_PROVIDER_URL,
            "model": requested_model or settings.AGENT_MODEL or "gpt-4o-mini",
            "source": "system_agent_provider",
            "name": "System Agent Provider",
        }
        _log(f"Using system Agent Provider (model={provider['model']})")
        return provider

    default_setting = (
        db.query(AISettings)
        .filter(AISettings.is_default == True, AISettings.is_active == True)
        .first()
    )
    if default_setting and default_setting.api_url and default_setting.api_key:
        provider = _ai_setting_provider(default_setting, requested_model, "ai_settings_default_provider")
        _log(f"Using AI Settings default provider '{provider['name']}' ({provider['provider']}, model={provider['model']})")
        return provider

    if settings.OPENAI_API_KEY:
        provider = {
            "provider": "openai_compatible",
            "apiKey": settings.OPENAI_API_KEY,
            "baseUrl": None,
            "model": requested_model or "gpt-4o-mini",
            "source": "system_openai_api_key",
            "name": "OPENAI_API_KEY",
        }
        _log(f"Using system OPENAI_API_KEY (model={provider['model']})")
        return provider

    fallback = (
        db.query(Integration)
        .filter(Integration.type == IntegrationType.LLM, Integration.status == IntegrationStatus.ACTIVE)
        .order_by(Integration.created_at.asc())
        .first()
    )
    if fallback:
        icfg = fallback.config or {}
        api_key = icfg.get("apiKey")
        if api_key:
            provider = {
                "provider": "openai_compatible",
                "apiKey": api_key,
                "baseUrl": icfg.get("baseUrl") or None,
                "model": requested_model or icfg.get("model") or "gpt-4o-mini",
                "source": "fallback_llm_integration",
                "name": fallback.name,
            }
            _log(f"Using fallback LLM integration '{fallback.name}' (model={provider['model']})")
            return provider

    raise NodeExecutionError("ยังไม่ได้ตั้งค่า LLM — ตั้งค่า Setting AI > Agent Provider หรือ Default AI provider ก่อน")


def _call_openai_compatible(provider: Dict[str, Any], messages: List[Dict[str, str]]) -> str:
    from openai import OpenAI

    client_kwargs: Dict[str, Any] = {"api_key": provider["apiKey"]}
    base_url = _normalize_openai_base_url(provider.get("baseUrl"))
    if base_url:
        client_kwargs["base_url"] = base_url
    client = OpenAI(**client_kwargs)
    try:
        response = client.chat.completions.create(model=provider["model"], messages=messages, temperature=0.2)
    except Exception as exc:  # noqa: BLE001 - some providers reject temperature
        if "temperature" not in str(exc).lower():
            raise
        response = client.chat.completions.create(model=provider["model"], messages=messages)
    return response.choices[0].message.content or "" if response.choices else ""


def _call_completion_messages(provider: Dict[str, Any], prompt: str, system_prompt: str) -> str:
    payload = {
        "inputs": {
            "ocr_content": prompt,
            "document_type": "workflow",
            "workflow_prompt": prompt,
            "agent_prompt": prompt,
            "system_prompt": system_prompt,
            "model": provider.get("model"),
        },
        "user": "workflow",
        "citation": False,
        "response_mode": "blocking",
    }
    response = http_requests.post(
        provider["apiUrl"],
        json=payload,
        headers={
            "Authorization": f"Bearer {provider['apiKey']}",
            "Content-Type": "application/json",
        },
        timeout=120,
    )
    response.raise_for_status()
    try:
        data = response.json()
    except ValueError:
        return response.text
    return str(data.get("answer") or data.get("text") or data.get("output") or data.get("result") or "")


def call_llm_provider(provider: Dict[str, Any], prompt: str, system_prompt: Optional[str] = None) -> str:
    system_text = (system_prompt or "").strip()
    if provider.get("provider") == "completion_messages":
        return _call_completion_messages(provider, prompt, system_text)

    messages: List[Dict[str, str]] = []
    if system_text:
        messages.append({"role": "system", "content": system_text})
    messages.append({"role": "user", "content": prompt})
    return _call_openai_compatible(provider, messages)


# Backward-compatible helper for callers that still expect an OpenAI client.
def resolve_llm_client(
    db: Session,
    integration_id: Optional[str] = None,
    model: Optional[str] = None,
    log: Optional[Callable[[str], None]] = None,
):
    from openai import OpenAI

    provider = resolve_llm_provider(db, integration_id, model, log)
    if provider.get("provider") != "openai_compatible":
        raise NodeExecutionError("Selected AI provider is not OpenAI-compatible")
    client_kwargs: Dict[str, Any] = {"api_key": provider["apiKey"]}
    base_url = _normalize_openai_base_url(provider.get("baseUrl"))
    if base_url:
        client_kwargs["base_url"] = base_url
    return OpenAI(**client_kwargs), provider["model"]

def _exec_llm(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    prompt = config.get("prompt")
    if not prompt:
        raise NodeExecutionError("LLM node: prompt is required")

    provider = resolve_llm_provider(
        db,
        config.get("integration_id"),
        config.get("model"),
        log,
        config.get("ai_provider_id"),
        owner_user_id=context.get("_owner_user_id"),
    )
    text = call_llm_provider(provider, _stringify(prompt), config.get("system_prompt"))
    log(f"LLM responded via {provider.get('source')} ({len(text)} chars)")

    if config.get("json_output"):
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip())
        try:
            return {"text": text, "data": json.loads(cleaned)}
        except json.JSONDecodeError:
            log("Warning: LLM output is not valid JSON — returning raw text")
            return {"text": text, "data": None}
    return {"text": text}


def suggest_variables(
    db: Session,
    query: str,
    candidates: List[Dict[str, Any]],
    integration_id: Optional[str] = None,
    owner_user_id: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """AI variable finder. Given a natural-language description and a catalog
    of available variables (token/label/sample), ask the LLM to rank the best
    matches. Returns [{token, reason, confidence}] limited to known tokens.

    The LLM only *selects* from the supplied tokens — sample values are
    rendered client-side from real run data, so it can't fabricate values.
    """
    valid_tokens = {c.get("token") for c in candidates if c.get("token")}
    if not valid_tokens:
        return []

    # Compact catalog so the prompt stays small even with many fields.
    lines = []
    for c in candidates[:200]:
        token = c.get("token")
        if not token:
            continue
        label = str(c.get("label") or "").strip()
        sample = str(c.get("sample") or "").replace("\n", " ")[:80]
        ctype = c.get("type") or ""
        lines.append(f"- {token} | label: {label} | type: {ctype} | ตัวอย่าง: {sample}")
    catalog = "\n".join(lines)

    system = (
        "You are a data-field matcher for a no-code workflow builder. "
        "The user describes (in Thai or English) the data they want to insert. "
        "Choose the variable tokens from the provided catalog that best match the request. "
        "Match on meaning, label, sample values, and field naming — Thai and English are equivalent "
        "(e.g. 'เลขที่ใบแจ้งหนี้' ≈ 'invoice number' ≈ 'Invoice_No'). "
        "Return STRICT JSON only: an array of at most 5 objects "
        '{"token": <exact token string from the catalog>, '
        '"reason": <short Thai explanation>, '
        '"confidence": <"high"|"medium"|"low">}. '
        "Order by relevance, best first. Use ONLY tokens that appear verbatim in the catalog. "
        "If nothing matches, return an empty array []."
    )
    user = f"คำขอของผู้ใช้: {query}\n\nรายการตัวแปรที่มี (catalog):\n{catalog}"

    provider = resolve_llm_provider(db, integration_id, None, None, owner_user_id=owner_user_id)
    text = call_llm_provider(provider, user, system)
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip())

    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        # Some models wrap the array in an object — try to dig it out.
        m = re.search(r"\[.*\]", cleaned, re.DOTALL)
        if not m:
            raise NodeExecutionError("AI ไม่สามารถตีความผลลัพธ์ได้ ลองพิมพ์คำอธิบายใหม่")
        parsed = json.loads(m.group(0))

    if isinstance(parsed, dict):
        # Tolerate {"results": [...]} or {"matches": [...]}
        for key in ("results", "matches", "data", "tokens"):
            if isinstance(parsed.get(key), list):
                parsed = parsed[key]
                break
        else:
            parsed = [parsed]

    results: List[Dict[str, Any]] = []
    seen: set = set()
    for item in parsed if isinstance(parsed, list) else []:
        if not isinstance(item, dict):
            continue
        token = item.get("token")
        if token not in valid_tokens or token in seen:
            continue
        seen.add(token)
        conf = str(item.get("confidence") or "medium").lower()
        if conf not in ("high", "medium", "low"):
            conf = "medium"
        results.append({
            "token": token,
            "reason": str(item.get("reason") or "")[:200],
            "confidence": conf,
        })
        if len(results) >= 5:
            break
    return results


def _evaluate_condition(left: Any, operator: str, right: Any) -> bool:
    def as_number(v: Any) -> Optional[float]:
        try:
            return float(v)
        except (TypeError, ValueError):
            return None

    if operator == "is_empty":
        return left is None or left == "" or left == [] or left == {}
    if operator == "is_not_empty":
        return not (left is None or left == "" or left == [] or left == {})
    if operator == "contains":
        return _stringify(right) in _stringify(left)
    if operator == "not_contains":
        return _stringify(right) not in _stringify(left)
    if operator in ("greater_than", "less_than"):
        ln, rn = as_number(left), as_number(right)
        if ln is None or rn is None:
            raise NodeExecutionError(f"Condition: cannot compare non-numeric values ({left!r} vs {right!r})")
        return ln > rn if operator == "greater_than" else ln < rn
    if operator == "not_equals":
        return _stringify(left) != _stringify(right)
    return _stringify(left) == _stringify(right)


def _exec_condition(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    left = config.get("left")
    right = config.get("right")
    operator = config.get("operator") or "equals"
    result = _evaluate_condition(left, operator, right)

    log(f"Condition: {left!r} {operator} {right!r} → {result}")
    return {"result": result}


def _exec_transform(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    mappings = config.get("mappings") or []
    out: Dict[str, Any] = {}
    for m in mappings:
        target = (m or {}).get("target")
        if target:
            out[target] = m.get("value")
    log(f"Transform produced {len(out)} field(s)")
    return out


def _exec_python_code(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    from app.services.code_sandbox import execute_python

    code = config.get("code")
    if not code:
        raise NodeExecutionError("Python node: code is required")
    timeout = int(config.get("timeout") or 30)
    node_input = config.get("input")
    if isinstance(node_input, str) and node_input.strip():
        try:
            node_input = json.loads(node_input)
        except json.JSONDecodeError:
            node_input = {"value": node_input}
    inputs = node_input if isinstance(node_input, dict) else {"value": node_input}

    log("Executing Python code in sandbox…")
    result = asyncio.run(execute_python(code, inputs=inputs, timeout=timeout))
    stdout = result.get("stdout")
    if stdout:
        log(f"stdout:\n{stdout}")
    if result.get("error"):
        err = result["error"]
        message = err.get("message") if isinstance(err, dict) else str(err)
        raise NodeExecutionError(f"Python error: {message}")
    return {"result": result.get("result"), "stdout": stdout}


def _validate_outbound_url(url: str) -> None:
    """SSRF guard for user-configurable HTTP nodes: only http(s), and no
    private/loopback/link-local/metadata targets unless explicitly allowed
    via WORKFLOW_HTTP_ALLOW_PRIVATE=true (e.g. for on-prem internal APIs)."""
    import ipaddress
    import socket
    from urllib.parse import urlparse

    if os.environ.get("WORKFLOW_HTTP_ALLOW_PRIVATE", "").lower() in ("1", "true", "yes"):
        return

    parsed = urlparse(str(url))
    if parsed.scheme not in ("http", "https"):
        raise NodeExecutionError(f"HTTP node: unsupported URL scheme '{parsed.scheme}'")
    host = parsed.hostname
    if not host:
        raise NodeExecutionError("HTTP node: URL has no host")
    try:
        infos = socket.getaddrinfo(host, None)
    except socket.gaierror as exc:
        raise NodeExecutionError(f"HTTP node: cannot resolve host '{host}': {exc}")
    for info in infos:
        ip = ipaddress.ip_address(info[4][0])
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast:
            raise NodeExecutionError(
                f"HTTP node: target '{host}' resolves to a private address ({ip}) — blocked. "
                "Set WORKFLOW_HTTP_ALLOW_PRIVATE=true to allow internal targets."
            )


def _exec_http_request(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    url = config.get("url")
    if not url:
        raise NodeExecutionError("HTTP node: url is required")
    _validate_outbound_url(url)
    method = (config.get("method") or "POST").upper()

    headers: Dict[str, str] = {}
    raw_headers = config.get("headers")
    if isinstance(raw_headers, dict):
        headers = {str(k): str(v) for k, v in raw_headers.items()}
    elif isinstance(raw_headers, str) and raw_headers.strip():
        try:
            headers = json.loads(raw_headers)
        except json.JSONDecodeError:
            raise NodeExecutionError("HTTP node: headers must be valid JSON")

    body = config.get("body")
    kwargs: Dict[str, Any] = {"headers": headers, "timeout": 60}
    if body is not None and method != "GET":
        if isinstance(body, (dict, list)):
            kwargs["json"] = body
        else:
            try:
                kwargs["json"] = json.loads(str(body))
            except json.JSONDecodeError:
                kwargs["data"] = str(body)

    log(f"{method} {url}")
    resp = http_requests.request(method, url, **kwargs)
    log(f"Response: HTTP {resp.status_code}")
    try:
        payload: Any = resp.json()
    except ValueError:
        payload = resp.text[:5000]
    if resp.status_code >= 400:
        raise NodeExecutionError(f"HTTP {resp.status_code}: {_stringify(payload)[:500]}")
    return {"status_code": resp.status_code, "body": payload}


_WRITE_OUTPUT_CONTENT_TYPES = {
    "json": "application/json",
    "csv": "text/csv",
    "text": "text/plain",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _exec_write_output(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    content = config.get("content")
    fmt = (config.get("format") or "json").lower()
    filename = os.path.basename(str(config.get("filename") or f"output.{fmt}"))  # no path traversal
    # Ensure the extension matches the format so the file opens correctly.
    if not filename.lower().endswith(f".{fmt}"):
        filename = f"{os.path.splitext(filename)[0] or 'output'}.{fmt}"

    run_id = context.get("_run_id", "unknown")

    text_preview: str
    if fmt in ("xlsx", "docx"):
        data = _to_xlsx_bytes(content) if fmt == "xlsx" else _to_docx_bytes(content)
        text_preview = f"({fmt.upper()} binary, {len(data)} bytes)"
    else:
        if fmt == "json":
            if isinstance(content, str):
                try:
                    content = json.loads(content)
                except json.JSONDecodeError:
                    pass
            text = json.dumps(content, ensure_ascii=False, indent=2)
        elif fmt == "csv":
            text = _to_csv(content)
        else:
            text = _stringify(content)
        data = text.encode("utf-8")
        text_preview = text[:2000]

    # Store via the storage service (shared volume / MinIO / S3) so the API
    # container can serve the file no matter which worker wrote it.
    from io import BytesIO
    from app.services.storage import get_storage_service

    storage_key = f"{WORKFLOW_OUTPUT_DIR}/{run_id}/{filename}"
    get_storage_service().upload_file(
        BytesIO(data),
        storage_key,
        content_type=_WRITE_OUTPUT_CONTENT_TYPES.get(fmt, "application/octet-stream"),
    )
    log(f"Wrote {len(data)} bytes to {storage_key}")
    return {"file_path": storage_key, "filename": filename, "size": len(data),
            "preview": text_preview, "run_id": str(run_id)}


def _exec_webhook_response(db: Session, config: dict, context: dict, log: Callable[[str], None]) -> Any:
    visible = bool(config.get("visible", True))
    operator = config.get("condition_operator") or ""
    if operator:
        left = config.get("condition_left")
        right = config.get("condition_right")
        visible = visible and _evaluate_condition(left, operator, right)
        log(f"Webhook response condition: {left!r} {operator} {right!r} → {visible}")

    status_code = int(config.get("status_code") or 200)
    if status_code < 100 or status_code > 599:
        raise NodeExecutionError("Webhook Response: status_code must be between 100 and 599")

    body = config.get("body")
    log("Webhook response prepared" if visible else "Webhook response hidden by condition")
    return {"visible": visible, "status_code": status_code, "body": body}


# ── Cloud storage (Google Drive / OneDrive) ──────────────────────────
def _load_drive_integration(db: Session, config: dict, expected_type: str, context: Optional[dict] = None):
    from app.services.cloud_drive import get_drive_client

    integration_id = config.get("integration_id")
    if not integration_id:
        raise NodeExecutionError("ต้องเลือก integration (บัญชีคลาวด์) ก่อน")
    integration = db.query(Integration).filter(Integration.id == integration_id).first()
    if not integration:
        raise NodeExecutionError(f"ไม่พบ integration: {integration_id}")
    _ensure_integration_owner(integration, (context or {}).get("_owner_user_id"))
    if integration.type != expected_type:
        raise NodeExecutionError(
            f"integration '{integration.name}' เป็นชนิด {integration.type} ไม่ใช่ {expected_type}"
        )
    return integration, get_drive_client(integration)


def _content_to_bytes(content: Any, mime_type: str) -> bytes:
    """Serialise node content to bytes — JSON object/array → pretty JSON, else str."""
    if isinstance(content, (dict, list)):
        return json.dumps(content, ensure_ascii=False, indent=2).encode("utf-8")
    if content is None:
        return b""
    if "json" in (mime_type or "").lower() and isinstance(content, str):
        # leave already-serialised JSON strings as-is
        return content.encode("utf-8")
    return _stringify(content).encode("utf-8")


def _exec_cloud_upload(db: Session, config: dict, context: dict, log, provider: str) -> Any:
    integration, client = _load_drive_integration(db, config, provider, context)
    filename = os.path.basename(str(config.get("filename") or "result.json"))
    mime_type = config.get("mime_type") or "application/json"
    folder_id = config.get("folder_id") or ""
    data = _content_to_bytes(config.get("content"), mime_type)

    log(f"อัปโหลด '{filename}' ({len(data)} bytes) ผ่าน '{integration.name}'")
    result = client.upload(folder_id, filename, data, mime_type)
    log(f"อัปโหลดสำเร็จ: {result.get('name')} (id={result.get('file_id')})")
    return result


def _exec_cloud_import(db: Session, config: dict, context: dict, log, provider: str) -> Any:
    from app.services.ingestion import (
        ingest_file_into_job,
        validate_import_file,
        DuplicateSourceFile,
        UnsupportedFile,
    )

    integration, client = _load_drive_integration(db, config, provider, context)
    folder_id = config.get("folder_id") or ""
    job_id = config.get("job_id")
    if not job_id:
        raise NodeExecutionError("ต้องเลือก Job ปลายทาง")
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise NodeExecutionError(f"ไม่พบ Job: {job_id}")

    name_filter = (config.get("name_filter") or "").strip().lower()
    limit = int(config.get("limit") or 20)

    files = client.list_folder(folder_id)
    if name_filter:
        files = [f for f in files if name_filter in (f.get("name") or "").lower()]

    # Dedup against files already imported into this job so a repeated /
    # scheduled import doesn't re-ingest (and re-pay for OCR of) the same files.
    already = {
        row[0]
        for row in db.query(Document.source_file_id)
        .filter(Document.job_id == job_id, Document.source_file_id.isnot(None))
        .all()
    }
    pending = [f for f in files if f.get("id") not in already]
    skipped_existing = len(files) - len(pending)
    pending = pending[:limit]
    log(
        f"พบ {len(files)} ไฟล์ในโฟลเดอร์ (นำเข้าแล้ว {skipped_existing}) — "
        f"เริ่มนำเข้า {len(pending)} ไฟล์ใหม่เข้า Job '{job.name or job_id}'"
    )

    imported: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []
    failed: List[Dict[str, Any]] = []
    for f in pending:
        fid, fname = f.get("id"), f.get("name") or f.get("id")
        try:
            # Pre-download guard using source metadata size: reject an oversize
            # or unsupported file BEFORE pulling its bytes into memory, so a
            # stray huge file in the folder can't OOM the worker.
            raw_size = f.get("size")
            try:
                meta_size = int(raw_size) if raw_size is not None else None
            except (TypeError, ValueError):
                meta_size = None
            validate_import_file(fname, meta_size)

            data = client.download(fid)
            res = ingest_file_into_job(
                db, str(job_id), data, fname, f.get("mimeType"), source_file_id=fid
            )
            imported.append({"document_id": res["document_id"], "filename": fname, "drive_file_id": fid})
            log(f"นำเข้า '{fname}' → document {res['document_id']}")
        except UnsupportedFile as exc:
            failed.append({"filename": fname, "drive_file_id": fid, "error": str(exc)})
            log(f"ข้าม '{fname}': {exc}")
        except DuplicateSourceFile:
            # Won a concurrent race; another run already imported this file.
            skipped.append({"filename": fname, "drive_file_id": fid})
            log(f"ข้าม '{fname}': นำเข้าแล้ว (ซ้ำ)")
        except Exception as exc:  # noqa: BLE001 — record a bad file, keep importing
            db.rollback()
            failed.append({"filename": fname, "drive_file_id": fid, "error": str(exc)[:500]})
            log(f"ข้าม '{fname}': {exc}")
    return {
        "job_id": str(job_id),
        "count": len(imported),
        "imported": imported,
        "skipped_count": skipped_existing + len(skipped),
        "skipped": skipped,
        "failed_count": len(failed),
        "failed": failed,
    }


def _exec_gdrive_upload(db, config, context, log):
    return _exec_cloud_upload(db, config, context, log, "gdrive")


def _exec_gdrive_import(db, config, context, log):
    return _exec_cloud_import(db, config, context, log, "gdrive")


def _exec_onedrive_upload(db, config, context, log):
    return _exec_cloud_upload(db, config, context, log, "onedrive")


def _exec_onedrive_import(db, config, context, log):
    return _exec_cloud_import(db, config, context, log, "onedrive")


def _tabular_rows(content: Any) -> Optional[List[dict]]:
    """Normalise content to a list-of-dicts table, or None if not tabular.

    Accepts a JSON string, a single dict (→ one row), or a list of dicts.
    Shared by the csv / xlsx / docx writers.
    """
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except json.JSONDecodeError:
            return None
    if isinstance(content, dict):
        return [content]
    if isinstance(content, list) and content and all(isinstance(r, dict) for r in content):
        return content
    return None


def _fieldnames(rows: List[dict]) -> List[str]:
    names: List[str] = []
    for r in rows:
        for k in r.keys():
            if k not in names:
                names.append(k)
    return names


def _cell(value: Any) -> str:
    return _stringify(value) if isinstance(value, (dict, list)) else ("" if value is None else str(value))


def _to_csv(content: Any) -> str:
    import csv
    import io

    rows = _tabular_rows(content)
    if rows is None:
        return content if isinstance(content, str) else _stringify(content)

    fieldnames = _fieldnames(rows)
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=fieldnames, extrasaction="ignore")
    writer.writeheader()
    for r in rows:
        writer.writerow({k: _cell(v) for k, v in r.items()})
    return buf.getvalue()


def _to_xlsx_bytes(content: Any) -> bytes:
    """Render content as an .xlsx workbook. Tabular content → header + rows;
    non-tabular → the stringified value in a single cell."""
    import io
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Output"
    rows = _tabular_rows(content)
    if rows is not None:
        fieldnames = _fieldnames(rows)
        ws.append(fieldnames)
        for r in rows:
            ws.append([_cell(r.get(k)) for k in fieldnames])
    else:
        ws["A1"] = content if isinstance(content, str) else _stringify(content)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _to_docx_bytes(content: Any) -> bytes:
    """Render content as a .docx document. Tabular content → a table; otherwise
    the stringified value as paragraphs (one per line)."""
    import io
    from docx import Document as DocxDocument

    doc = DocxDocument()
    rows = _tabular_rows(content)
    if rows is not None:
        fieldnames = _fieldnames(rows)
        table = doc.add_table(rows=1, cols=len(fieldnames) or 1)
        try:
            table.style = "Table Grid"
        except Exception:  # style may be unavailable in a minimal template
            pass
        hdr = table.rows[0].cells
        for i, name in enumerate(fieldnames):
            hdr[i].text = str(name)
        for r in rows:
            cells = table.add_row().cells
            for i, name in enumerate(fieldnames):
                cells[i].text = _cell(r.get(name))
    else:
        text = content if isinstance(content, str) else _stringify(content)
        for line in text.split("\n"):
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


EXECUTORS: Dict[str, Callable] = {
    "trigger_manual": _exec_trigger,
    "trigger_schedule": _exec_trigger,
    "trigger_webhook": _exec_trigger,
    "job_source": _exec_job_source,
    "document_source": _exec_document_source,
    "llm": _exec_llm,
    "condition": _exec_condition,
    "transform": _exec_transform,
    "python_code": _exec_python_code,
    "http_request": _exec_http_request,
    "write_output": _exec_write_output,
    "webhook_response": _exec_webhook_response,
    "gdrive_upload": _exec_gdrive_upload,
    "gdrive_import": _exec_gdrive_import,
    "onedrive_upload": _exec_onedrive_upload,
    "onedrive_import": _exec_onedrive_import,
}


# ── Engine ───────────────────────────────────────────────────────────
def _now() -> datetime:
    return datetime.now(timezone.utc)


def _topological_order(nodes: List[dict], edges: List[dict]) -> List[dict]:
    node_map = {n["id"]: n for n in nodes}
    indegree = {nid: 0 for nid in node_map}
    children: Dict[str, List[str]] = {nid: [] for nid in node_map}
    for e in edges:
        src, dst = e.get("source"), e.get("target")
        if src in node_map and dst in node_map:
            indegree[dst] += 1
            children[src].append(dst)

    queue = [nid for nid, deg in indegree.items() if deg == 0]
    ordered: List[dict] = []
    while queue:
        nid = queue.pop(0)
        ordered.append(node_map[nid])
        for child in children[nid]:
            indegree[child] -= 1
            if indegree[child] == 0:
                queue.append(child)
    if len(ordered) != len(nodes):
        raise NodeExecutionError("Workflow contains a cycle — must be a DAG")
    return ordered


def _workflow_owner_id(db: Session, workflow_id) -> Optional[str]:
    # Best-effort: on any failure (orphaned run, non-query session), return
    # None, which disables cross-tenant enforcement rather than killing the
    # run. Real runs use a full Session and resolve the owner normally.
    try:
        owner = (
            db.query(Workflow.user_id)
            .filter(Workflow.id == workflow_id)
            .scalar()
        )
    except Exception:
        return None
    return str(owner) if owner else None


def execute_workflow_run(db: Session, run: WorkflowRun) -> None:
    """Execute one workflow run synchronously, persisting node activity."""
    definition = run.definition_snapshot or {}
    nodes: List[dict] = definition.get("nodes") or []
    edges: List[dict] = definition.get("edges") or []

    run.status = "running"
    run.started_at = _now()
    db.commit()

    # Pre-create node-run rows so the UI immediately shows pending steps
    node_run_map: Dict[str, WorkflowNodeRun] = {}
    for n in nodes:
        nr = WorkflowNodeRun(
            run_id=run.id,
            node_id=n["id"],
            node_type=n.get("type", "unknown"),
            node_label=(n.get("data") or {}).get("label"),
            status="pending",
        )
        db.add(nr)
        node_run_map[n["id"]] = nr
    db.commit()

    context: Dict[str, Any] = {
        "trigger": run.trigger_input or {},
        "_run_id": str(run.id),
        "_owner_user_id": _workflow_owner_id(db, run.workflow_id),
    }

    incoming: Dict[str, List[dict]] = {}
    for e in edges:
        incoming.setdefault(e.get("target"), []).append(e)

    trigger_types = {"trigger_manual", "trigger_schedule", "trigger_webhook"}
    node_status: Dict[str, str] = {}
    run_failed_error: Optional[str] = None
    fallback_result: Any = None
    fallback_result_node_id: Optional[str] = None
    webhook_result: Any = None
    webhook_result_node_id: Optional[str] = None

    try:
        ordered = _topological_order(nodes, edges)
    except NodeExecutionError as exc:
        run.status = "failed"
        run.error = str(exc)
        run.finished_at = _now()
        db.commit()
        return

    for node in ordered:
        node_id = node["id"]
        node_type = node.get("type", "unknown")
        nr = node_run_map[node_id]

        # Decide whether this node should execute
        should_run = False
        if node_type in trigger_types:
            should_run = not incoming.get(node_id)
        in_edges = incoming.get(node_id) or []
        if in_edges and run_failed_error is None:
            for e in in_edges:
                src = e.get("source")
                if node_status.get(src) != "succeeded":
                    continue
                src_output = context.get(src)
                src_node = next((n for n in nodes if n["id"] == src), None)
                if src_node and src_node.get("type") == "condition":
                    branch = str((src_output or {}).get("result", False)).lower()
                    handle = (e.get("sourceHandle") or "true").lower()
                    if handle == branch:
                        should_run = True
                        break
                else:
                    should_run = True
                    break

        if run_failed_error is not None or not should_run:
            nr.status = "skipped"
            nr.finished_at = _now()
            node_status[node_id] = "skipped"
            db.commit()
            continue

        # Execute the node
        logs: List[str] = []

        def log(msg: str) -> None:
            timestamp = _now().strftime("%H:%M:%S")
            logs.append(f"[{timestamp}] {msg}")

        raw_config = (node.get("data") or {}).get("config") or {}
        nr.status = "running"
        nr.started_at = _now()
        db.commit()

        try:
            resolved_config = resolve_template(raw_config, context)
            nr.input = _safe_json(redact_secrets(resolved_config))
            executor = EXECUTORS.get(node_type)
            if not executor:
                raise NodeExecutionError(f"Unknown node type: {node_type}")
            output = executor(db, resolved_config, context, log)
            context[node_id] = output
            nr.output = _safe_json(output)
            nr.status = "succeeded"
            node_status[node_id] = "succeeded"
            if node_type == "webhook_response":
                if isinstance(output, dict) and output.get("visible"):
                    webhook_result = output
                    webhook_result_node_id = node_id
            else:
                fallback_result = output
                fallback_result_node_id = node_id
        except Exception as exc:  # noqa: BLE001 — node failures must not kill the loop
            logger.exception("Workflow node %s failed", node_id)
            nr.status = "failed"
            nr.error = str(exc)[:4000]
            node_status[node_id] = "failed"
            run_failed_error = f"Node '{(node.get('data') or {}).get('label') or node_id}' failed: {exc}"
        finally:
            nr.logs = "\n".join(logs) if logs else None
            nr.finished_at = _now()
            db.commit()

    run.status = "failed" if run_failed_error else "succeeded"
    run.error = run_failed_error
    if not run_failed_error:
        selected_result = webhook_result if webhook_result is not None else fallback_result
        run.result = _safe_json(selected_result)
        run.result_node_id = webhook_result_node_id or fallback_result_node_id
    run.finished_at = _now()
    db.commit()


def _build_context_from_last_run(db: Session, workflow_id: Any) -> Optional[Dict[str, Any]]:
    """Reconstruct an execution context from the most recent full run of a workflow.

    Returns {node_id: output, ..., "trigger": <trigger_input>} or None if there is
    no prior full run (manual/schedule) to borrow data from.
    """
    last_run = (
        db.query(WorkflowRun)
        .filter(
            WorkflowRun.workflow_id == workflow_id,
            WorkflowRun.trigger_type.in_(["manual", "schedule"]),
        )
        .order_by(WorkflowRun.created_at.desc())
        .first()
    )
    if not last_run:
        return None

    context: Dict[str, Any] = {"trigger": last_run.trigger_input or {}}
    for nr in last_run.node_runs:
        if nr.output is not None:
            context[nr.node_id] = nr.output
    return context


def execute_single_node(db: Session, run: WorkflowRun, node_id: str) -> None:
    """Execute one node in isolation, borrowing upstream data from the last full run.

    Persists a single WorkflowNodeRun so the existing Activity panel can show it.
    """
    definition = run.definition_snapshot or {}
    nodes: List[dict] = definition.get("nodes") or []
    node = next((n for n in nodes if n["id"] == node_id), None)

    run.status = "running"
    run.started_at = _now()
    db.commit()

    if not node:
        run.status = "failed"
        run.error = f"ไม่พบโหนด {node_id} ใน workflow"
        run.finished_at = _now()
        db.commit()
        return

    node_type = node.get("type", "unknown")
    nr = WorkflowNodeRun(
        run_id=run.id,
        node_id=node_id,
        node_type=node_type,
        node_label=(node.get("data") or {}).get("label"),
        status="pending",
    )
    db.add(nr)
    db.commit()

    # Trigger nodes have no upstream — just echo the last run's trigger input
    is_trigger = node_type in ("trigger_manual", "trigger_schedule", "trigger_webhook")
    context = _build_context_from_last_run(db, run.workflow_id)
    if context is None and not is_trigger:
        nr.status = "failed"
        nr.error = "กรุณารัน workflow แบบเต็มอย่างน้อย 1 ครั้งก่อน เพื่อให้มีข้อมูลจากโหนดก่อนหน้า"
        nr.finished_at = _now()
        db.commit()
        run.status = "failed"
        run.error = nr.error
        run.finished_at = _now()
        db.commit()
        return
    if context is None:
        context = {"trigger": {}}
    context["_run_id"] = str(run.id)
    context["_owner_user_id"] = _workflow_owner_id(db, run.workflow_id)

    logs: List[str] = []

    def log(msg: str) -> None:
        logs.append(f"[{_now().strftime('%H:%M:%S')}] {msg}")

    log("ทดสอบโหนดนี้ด้วยข้อมูลจากการรันเต็มครั้งล่าสุด")
    raw_config = (node.get("data") or {}).get("config") or {}
    nr.status = "running"
    nr.started_at = _now()
    db.commit()

    error: Optional[str] = None
    try:
        resolved_config = resolve_template(raw_config, context)
        nr.input = _safe_json(redact_secrets(resolved_config))
        executor = EXECUTORS.get(node_type)
        if not executor:
            raise NodeExecutionError(f"Unknown node type: {node_type}")
        output = executor(db, resolved_config, context, log)
        nr.output = _safe_json(output)
        nr.status = "succeeded"
    except Exception as exc:  # noqa: BLE001
        logger.exception("Single-node test %s failed", node_id)
        nr.status = "failed"
        nr.error = str(exc)[:4000]
        error = f"โหนด '{(node.get('data') or {}).get('label') or node_id}' ล้มเหลว: {exc}"
    finally:
        nr.logs = "\n".join(logs) if logs else None
        nr.finished_at = _now()
        db.commit()

    run.status = "failed" if error else "succeeded"
    run.error = error
    run.finished_at = _now()
    db.commit()


def _safe_json(value: Any) -> Any:
    """Ensure a value is JSON-serializable (truncate huge strings)."""
    try:
        text = json.dumps(value, ensure_ascii=False, default=str)
    except (TypeError, ValueError):
        return {"repr": repr(value)[:2000]}
    if len(text) > 200_000:
        return {"truncated": True, "preview": text[:10_000]}
    return json.loads(text)
